import csv
import os
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from config_class import Config
from header_numerizer import HeaderNumerizer


class DocxGenerator:
    def __init__(self):
        self.__document: Document = None
        self.__style = None
        self.__font = None
        self.config = Config("config.ini")
        self.__file_path = self.config.output_docx_path
        self.__create_file()
        self.function_map = {}
        self.__create_function_map()
        self.__header_numerizer = HeaderNumerizer()

    def __create_file(self):
        if not os.path.exists(self.__file_path):
            self.__document = Document()
        else:
            self.__document = Document(self.__file_path)

        self.__style = self.__document.styles[self.config.document_style]
        self.__font = self.__style.font
        self.save_document()

    def __create_function_map(self):
        self.function_map = {
            name: getattr(self, name)
            for name in dir(self)
            if callable(getattr(self, name)) and not name.startswith('__')
        }

    def set_document_style(self, style):
        self.__style = self.__document.styles[style]
        self.save_document()

    def set_document_font(self, font, size):
        self.__font.name = font
        self.__font.size = Pt(size)
        self.save_document()

    def add_page_break(self):
        """
        Add a page break to the document, starting a new page.
        """
        self.__document.add_page_break()
        self.save_document()

    def add_text(self, text, bold=False, italic=False):
        p = self.__document.add_paragraph("")
        if bold:
            p.add_run(text).bold = True
        elif italic:
            p.add_run(text).italic = True
        else:
            self.__document.add_paragraph(text)
        self.save_document()

    def add_bullet_point(self, text, numbers=False):
        if numbers:
            self.__document.add_paragraph(text, style='List Number')
        else:
            self.__document.add_paragraph(text, style='List Bullet')
        self.save_document()

    def add_heading(self, text, level=1):
        text = self.__header_numerizer.add_heading(text, level)
        self.__document.add_heading(text, level=level)
        self.save_document()

    def add_figure(self, image_path, title, width=None, height=None):
        """
        Add a figure (image) to the document.

        Parameters:
            image_path (str): The path to the image file.
            title (str): The title for the figure. Default is None.
            width (float): The width of the image in inches. Default is None.
            height (float): The height of the image in inches. Default is None.
        """
        self.__document.add_paragraph(title, style='Caption')
        if width is None and height is None:
            self.__document.add_picture(image_path)
        elif width is not None and height is not None:
            self.__document.add_picture(image_path, width=Inches(width), height=Inches(height))
        elif width is not None:
            self.__document.add_picture(image_path, width=Inches(width))
        elif height is not None:
            self.__document.add_picture(image_path, height=Inches(height))
        self.save_document()

    def add_table(self, csv_path, heading=None):
        """
        Add a table to the document.

        Parameters:
            rows (int): The number of rows in the table.
            cols (int): The number of columns in the table.

        Returns:
            Table: The added table object.
            :param heading:
            :param csv_path:
        """
        if heading:
            self.add_heading(heading)
        csv_table = self.__load_table_from_csv(csv_path)
        rows = len(csv_table)
        cols = len(csv_table[0])
        table = self.__document.add_table(rows, cols)
        table.style = 'Table Grid'
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = csv_table[i][j]
        self.save_document()

    def add_list_of_figures(self, list_of_figures):
        self.add_page_break()
        self.add_heading("list of figures:", 0)
        for figure_title in list_of_figures:
            self.add_text(figure_title, bold=True)
        self.add_page_break()

    def save_document(self):
        self.__document.save(self.__file_path)

    @staticmethod
    def __load_table_from_csv(file_path):
        """
        Load a table from a CSV file.

        Parameters:
            file_path (str): The path to the CSV file.

        Returns:
            list: A list of lists representing the table.
        """
        table = []
        with open(file_path, 'r', newline='') as csvfile:
            csv_reader = csv.reader(csvfile)
            for row in csv_reader:
                table.append(row)
        return table


# Example usage:


def main():
    # Create a WordDocument instance
    word_doc = DocxGenerator()

    # Add a heading
    word_doc.function_map["add_heading"]('Heading level 0', level=0)
    word_doc.function_map["add_heading"]('Heading level 1', level=1)
    word_doc.add_text('This is another paragraph.')
    word_doc.function_map["add_heading"]('Heading level 1', level=1)
    word_doc.add_text('This is another paragraph.')
    word_doc.function_map["add_heading"]('Heading level 2', level=2)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.function_map["add_heading"]('Heading level 2', level=2)
    word_doc.function_map["add_heading"]('Heading level 2', level=2)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.function_map["add_heading"]('Heading level 2', level=2)
    word_doc.function_map["add_heading"]('Heading level 0', level=0)
    word_doc.function_map["add_heading"]('Heading level 1', level=1)
    word_doc.add_text('This is another paragraph.')
    word_doc.function_map["add_heading"]('Heading level 1', level=1)
    word_doc.add_text('This is another paragraph.')
    word_doc.function_map["add_heading"]('Heading level 2', level=2)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.function_map["add_heading"]('Heading level 2', level=2)
    word_doc.function_map["add_heading"]('Heading level 2', level=2)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.add_figure(r'C:\Users\talgab\OneDrive - Mobileye\Pictures\Picture1.jpg', "Picture1.jpg", width=5, height=4)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.function_map["add_heading"]('Heading level 2', level=2)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.add_figure(r'C:\Users\talgab\OneDrive - Mobileye\Pictures\Picture1.jpg', "Picture1.jpg", width=5, height=4)
    word_doc.function_map["add_heading"]('Heading level 3', level=3)
    word_doc.add_figure(r'C:\Users\talgab\OneDrive - Mobileye\Pictures\Picture1.jpg', "Picture1.jpg", width=5, height=4)
    word_doc.function_map["add_heading"]('Heading level 3', level=4)
    word_doc.function_map["add_heading"]('Heading level 3', level=4)
    # Add table
    word_doc.add_table("datadata.csv")
    # Add more general paragraph text
    word_doc.add_text('This is another paragraph.')

    # Add bullet points with numbers
    word_doc.add_bullet_point('Numbered bullet point 1', numbers=True)
    word_doc.add_bullet_point('Numbered bullet point 2', numbers=True)
    word_doc.add_bullet_point('Numbered bullet point 3', numbers=True)

    # Add a figure without a title
    word_doc.add_figure(r'C:\Users\talgab\OneDrive - Mobileye\Pictures\Picture1.jpg',"Picture1.jpg", width=5, height=4)

    # Add more general paragraph text
    word_doc.add_text('This is the final paragraph.')


if __name__ == "__main__":
    main()
